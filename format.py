#file that given parameters
#format.py myDirectory
#takes the jpgs in that directory and outputs as pdfin current dir

import sys
import os
from docx import Document
from docx.shared import Inches

margin = 0.2

document = Document()

sections = document.sections
for section in sections:
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

path = './downloads/' + sys.argv[1]

p = document.add_paragraph()
r = p.add_run()
for image in os.listdir(path):
	print(image)
	r.add_text(' ')
	r.add_picture(path + "/" + image, width=Inches(2.5), height=Inches(3.5))

document.save('demo.docx')